from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()
#profile pic
document.add_picture('Screenshot_20230430_174741.png', width= Inches(2.0))
    


#name ph num and eamil details
name=input('what is ur name!')
speak('hello'+ name )
phone_number=input('what is ur ph no!')
email=input('what is ur email!')
document.add_paragraph(name+' | ' +phone_number+' | '+email)

#about me
document.add_heading('About me')
about_me=input('tell about your self')
document.add_paragraph(about_me)

#experiences
document.add_heading('work experiences')
p=document.add_paragraph()
company = input('enter company')
from_date=input('from date')
to_date=input('to_date')
experience_details=input('tell your experience at ' + company)
p.add_run(company+' ').bold=True
p.add_run(from_date+' - '+ to_date+'\n').italic=True
p.add_run(experience_details)

#has more experiences
while True:
    has_more_experiences=input('do you have more experiences?yes or no')
    if has_more_experiences.lower() == 'yes':
        p=document.add_paragraph()
        company = input('enter company')
        from_date=input('from date')
        to_date=input('to_date')
        experience_details=input('tell your experience at ' + company)
        p.add_run(company+' ').bold=True
        p.add_run(from_date+' - '+ to_date+'\n').italic=True
        p.add_run(experience_details)
    else:
        break   

#skills
document.add_heading('Skills')
skill = input('enter the skill')
p=document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills=input('do you have more skills? yes or no')
    if has_more_skills.lower()=='yes':
       skill = input('enter the skill')
       p=document.add_paragraph(skill)
       p.style = 'List Bullet'
    else:
        break


document.save('cv.docx')
